#!/usr/bin/env python
# -*- coding: latin-1 -*-
"""
    Read in a text file and create an XLSX spreadsheet.

    Assumes that the first line of the file that contains a
    comma is a header line.

    Strips off leading and ending "'" delimiters for items in each line,
    if they exist.
"""

from optparse import OptionParser
from collections import OrderedDict
import operator
import re
import sys
import os
import logging
import copy
import glob
import time
from difflib import Differ
from constants import *
from create_translation import *
from docx import Document
from docx.shared import Inches
import codecs
import subprocess
import datetime

class bit_field(object):
    def __init__(self):
        self.bit_range = ""
        self.bit_name = ""
        self.spec_part = ""
        self.spec_section = ""

class register_summary(object):
    def __init__(self):
        self.block = ""
        self.offset = ""
        self.name = ""
        self.bitlist = list(range(32))
        self.bits = []
        self.next_bit = 0;
        self.new_last_bit = -1;
    
class WordEditor(object):
    def __init__(self, text, word):
        self.text_filepath = text
        self.word_filepath = word
        self.regs = []
        self.lines = []
        self._read_text()
        self.doc = None

    def print_regs(self):
        for reg in self.regs:
            print ('\n%s' % "', '".join([reg.block, reg.offset, reg.name]))
            for bit in reg.bits:
                print ('%s' % "', '".join(
                [bit.bit_range, bit.bit_name, bit.spec_part, bit.spec_section]))

    def add_reserved_row(self, table, first_bit, last_bit):
        rsvd_cells = table.add_row().cells
        rsvd_cells[1].merge(rsvd_cells[3])
        bit_lable = ""
        if first_bit == last_bit:
            bit_lable = str(first_bit)
        else:
            bit_lable = "%d:%d" % (first_bit, last_bit)
        rsvd_cells[0].add_paragraph(bit_lable)
        rsvd_cells[1].add_paragraph("Reserved")

    def add_reserved_fields(self, table, prev_bit, bit, reg):
        new_bits = [tok.strip() for tok in bit.bit_range.split(":")]
        if len(new_bits) == 1:
            self.next_bit = int(new_bits[0])
            self.new_last_bit = self.next_bit + 1
        else:
            self.next_bit = int(new_bits[0])
            self.new_last_bit = int(new_bits[1]) + 1
        for bit_num in list(range(self.next_bit, self.new_last_bit)):
            if bit_num not in reg.bitlist:
                self.add_bit = False
        if ((self.next_bit > prev_bit) and self.add_bit):
            self.add_reserved_row(table, prev_bit, self.next_bit - 1)
        return self.new_last_bit

    def extract_part(self, part_string):
        start = part_string.find("Part ")
        if (start < 0):
            raise ValueError("Could not find part in %s!" % part_string);
        end = part_string[start:].find(":")
        if (end < 0):
            raise ValueError("Could not find end of part in %s!" % part_string);
        return part_string[start:start + end]

    def create_document(self):
        document = Document()

        document.add_heading('Registers Summary First Cut', 0)

        p = document.add_paragraph('Autogenerated register summary from file ')
        p.add_run(self.text_filepath).bold = True
        p.add_run('.  Generated ').bold = True
        now = datetime.datetime.now()
        p.add_run(now.strftime("%Y-%m-%d %H:%M:%S"))
        p = document.add_paragraph('Note that registers are defined using big '
                                   'endian notation.  Bit 0 is the most '
                                   'significant bit!')
        prev_block = "XXXX"
        table = None

        for reg in self.regs:
            if reg.block == prev_block:
                reg_cells = table.add_row().cells
            else:
                document.add_page_break()
                toks = [tok.strip() for tok in reg.name.split("Header")]
                document.add_heading('Block: %s : %s' % (reg.block, toks[0]),
                                      level=1)
                table = document.add_table(rows=1, cols=4)
                reg_cells = table.rows[0].cells
                prev_block = reg.block
            reg_cells[0].merge(reg_cells[3])
            reg_cells[0].add_paragraph("Name: %s\nOffset: %s" %
                                       (reg.name, reg.offset),
                                       style=None)
            hdr_cells = table.add_row().cells
            hdr_cells[0].text = 'Bits'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Part'
            hdr_cells[3].text = 'Section'

            last_bit = 0
            for bit in reg.bits:
                self.add_bit = True
                last_bit = self.add_reserved_fields(table, last_bit, bit, reg)
                if not self.add_bit:
                    continue
                for bit_num in list(range(self.next_bit, self.new_last_bit)):
                    reg.bitlist.remove(bit_num)
                row_cells = table.add_row().cells
                row_cells[0].text = bit.bit_range
                row_cells[1].text = bit.bit_name
                row_cells[2].text = self.extract_part(bit.spec_part)
                row_cells[3].text = bit.spec_section
            if not last_bit == 32:
                self.add_reserved_row(table, last_bit, 31)
        self.doc = document
    
    def _strip_line(self, line):
        toks = []
        if len(line):
            toks = [tok.strip() for tok in line[1:-1].split("', '")]
        return toks

    def _read_text(self):
        logging.info("Reading text file '%s'." % self.text_filepath)
        with open(self.text_filepath, 'r') as text_file:
            self.lines = [l.strip() for l in text_file.readlines()]
        col_warning = False
        reg = None

        for num, l in enumerate(self.lines):
            toks = self._strip_line(l)

            if len(toks) == 3:
                if reg is not None:
                    if not reg.block == "UNKNOWN":
                        self.regs.append(reg)
                reg = register_summary()
                reg.block = toks[0]
                reg.offset = toks[1]
                reg.name = toks[2]
                if (reg.offset[0] == '0'):
                    reg.offset = " " + reg.offset
            elif len(toks) == 4:
                if reg is None:
                    raise ValueError("Missing reg header at %d!" % num);
                bit = bit_field()
                bit.bit_range = toks[0]
                bit.bit_name = toks[1]
                bit.spec_part = toks[2]
                bit.spec_section = toks[3]
                reg.bits.append(bit)
            else:
                raise ValueError("Unknown line %d %s!" % (num, l));
        if reg is not None:
            if not reg.block == "UNKNOWN":
                self.regs.append(reg)

    def write_document(self, file_path):
        self.doc.save(file_path)

def create_parser():
    parser = OptionParser(description="Create Excel spreadsheet based on text file.")
    parser.add_option('-t', '--textfile',
            dest = 'text_filepath',
            action = 'store', type = 'string', default = None,
            help = 'File path to text file.',
            metavar = 'FILE')
    parser.add_option('-w', '--word',
            dest = 'word_filepath',
            action = 'store', type = 'string', default = None,
            help = 'File path to new Word file.',
            metavar = 'FILE')
    return parser

def validate_options(options):
    if not os.path.isfile(options.text_filepath):
        raise ValueError("Text file '%s' does not exist." %
                         options.text_filepath)
    #if os.path.isfile(options.word_filepath):
    #    raise ValueError("File '%s' will be overwritten!" %
    #                     options.word_filepath)

def main(argv = None):
    logging.basicConfig(level=logging.WARN)
    parser = create_parser()
    if argv is None:
        argv = sys.argv[1:]

    (options, argv) = parser.parse_args(argv)
    if len(argv) != 0:
        print('Invalid argument!')
        print
        parser.print_help()
        return -1

    try:
        validate_options(options)
    except ValueError as e:
        print(e)
        sys.exit(-1)

    word = WordEditor(options.text_filepath, options.word_filepath)
    word.create_document()
    word.write_document(options.word_filepath)

if __name__ == '__main__':
    sys.exit(main())
